import configparser
# log_config.py
import logging
import os

import requests
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fastmodels import Client

# 配置日志文件路径
LOG_DIR = '../logs'  # 日志目录
LOG_FILE = 'app.log'  # 日志文件名
LOG_PATH = os.path.join(LOG_DIR, LOG_FILE)

# 创建日志目录（如果不存在）
os.makedirs(LOG_DIR, exist_ok=True)

# 配置日志
logging.basicConfig(
    level=logging.DEBUG,  # 设置日志级别
    format='%(asctime)s - %(levelname)s - %(message)s',  # 日志格式
    handlers=[
        logging.FileHandler(LOG_PATH, encoding='utf-8'),
        logging.StreamHandler()  # 也可以输出到控制台

    ]
)

logger = logging.getLogger(__name__)  # 创建一个日志记录器


class Util:
    _instance = None

    def __init__(self):
        self.mail_domain = None
        self.mail_url = None
        self.mail_api_key = None
        self.oss_api_secret = None
        self.oss_api_key = None
        self.esm_project_id = None
        self.esm_id = None
        self.esm_api_key = None

    def __new__(self, *args, **kwargs):
        if not self._instance:
            self._instance = super(Util, self).__new__(self)
            self._instance.init()
        return self._instance

    def init(self):
        config_dict = configparser.ConfigParser()
        config_dict.read("../config.ini")
        self.oss_api_key = config_dict["OSS"]["id"]
        self.oss_api_secret = config_dict["OSS"]["secret"]
        self.mail_api_key = config_dict["MAILGUN"]["api_key"]
        self.mail_url = config_dict["MAILGUN"]["url"]
        self.mail_domain = config_dict["MAILGUN"]["domain"]
        logger.info("Mail URL: %s", self.mail_url)
        logger.info("Finished config init")

    def send_email(self, to_addresses, subject, context, file_path):
        response = requests.post(
            self.mail_url,
            auth=(self.mail_domain, self.mail_api_key),
            files=[("attachment", open(file_path, "rb"))],  # 添加附件
            data={"from": "operation@antpool.com",
                  "to": to_addresses,
                  "subject": subject,
                  "text": context})
        logger.info("Call mailgun %s", response.json())

def set_cell_border(cell, top=None, left=None, bottom=None, right=None):
    """
    设置单元格的边框。

    :param cell: 单元格对象
    :param top: 上边框宽度
    :param left: 左边框宽度
    :param bottom: 下边框宽度
    :param right: 右边框宽度
    """
    tc = cell._element  # 获取cell的xml元素

    # 使用正确的命名空间来查找 tcBorders
    namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # 获取单元格的边框对象，如果没有找到则创建新的
    tc_borders = tc.find('.//w:tcPr/w:tcBorders', namespaces=namespaces)
    
    if not tc_borders:
        # 如果没有边框，创建一个新的边框对象
        tc_pr = tc.get_or_add_tcPr()
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)

    # 设置边框
    def set_border(border_name, value):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')  # 设置边框类型为单线
        border.set(qn('w:sz'), str(value))  # 设置边框宽度
        border.set(qn('w:space'), '0')  # 设置间距
        return border

    # 设置每个边框
    if top is not None:
        tc_borders.append(set_border('top', top))
    if left is not None:
        tc_borders.append(set_border('left', left))
    if bottom is not None:
        tc_borders.append(set_border('bottom', bottom))
    if right is not None:
        tc_borders.append(set_border('right', right))

class MeetingNotesProcessor:
    def __init__(self, file_path):
        """
        初始化类，加载指定路径的 Word 文件。
        
        :param file_path: Word 文件的路径。
        """
        self.file_path = file_path
        self.document = Document(file_path)
        self.print_table_structure()  # 打印表格结构，调试时查看表格内容

    def print_table_structure(self):
        # 遍历文档中的每个表格
        for table in self.document.tables:
            print("Table:")
            # 遍历表格中的每一行
            for i, row in enumerate(table.rows):
                print(f"Row {i + 1}:")
                # 遍历每一行中的每个单元格
                for j, cell in enumerate(row.cells):
                    print(f"  Cell {j + 1}: {cell.text}")  # 打印单元格内容

    def save(self):
        """
        保存修改后的文档，并确保文件夹存在。
        """
        # 确保保存目录存在
        output_path = "result.docx"
        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)

        # 保存修改后的文件
        self.document.save(output_path)
        print(f"修改后的文档已保存为：{output_path}")

    def _record_content(self, keyword, content):
        """
        通用方法，用于记录内容到表格中。

        :param keyword: 查找的关键词
        :param content: 要插入的内容
        """
        for table in self.document.tables:
            for row in table.rows:
                # 遍历该行的所有单元格，检查关键词是否存在
                for idx, cell in enumerate(row.cells):
                    if keyword.strip() in cell.text.strip():  # 检查当前单元格是否包含关键词
                        if len(row.cells) > idx + 1:
                            # 如果该列后面有单元格，直接插入内容
                            row.cells[idx + 1].text = content
                        else:
                            # 如果该行没有足够的列，动态添加新单元格
                            row.add_cell()  # 向该行添加新单元格
                            row.cells[idx + 1].text = content  # 在新增单元格中插入内容
                        break  # 找到后就退出当前行的循环

    def add_action_item(self, serial_number, item, responsible, deadline):
        """
        添加一条行动计划。

        :param serial_number: 序号
        :param item: 事项
        :param responsible: 负责人
        :param deadline: 完成时间
        """
        action_item = [serial_number, item, responsible, deadline]
        
        # 查找表格中的“序号”标题所在位置
        for table in self.document.tables:
            for i, row in enumerate(table.rows):
                if '序号' in row.cells[0].text:  # 找到“序号”标题行
                    # 插入新的一行
                    new_row = table.add_row()  # 插入新行
                    
                    # 检查第一行是否为合并单元格
                    if len(table.rows[0].cells) == 1:  # 第一行只有一个单元格
                        new_row.cells[0].text = str(action_item[0])  # 填充序号
                    else:
                        # 填充新的行
                        new_row.cells[0].text = str(action_item[0])
                        new_row.cells[1].text = action_item[1]
                        new_row.cells[2].text = action_item[2]
                        new_row.cells[3].text = action_item[3]

                    # 确保第一列没有影响合并单元格
                    if len(table.rows[0].cells) == 1:
                        # 如果第一行是合并的，我们不填充第一列
                        new_row.cells[0].text = ""  # 保持第一列为空

                    # # 设置单元格边框
                    # for cell in new_row.cells:
                    #     set_cell_border(cell, top=4, left=4, bottom=4, right=4)  # 给单元格设置边框

                    break  # 找到后就退出循环

class AIClient:
    def __init__(self, api_key, project_id):
        """
        Initialize AI client and connect to the fastmodels API.
        
        :param api_key: API key for authentication.
        :param project_id: Project ID for the fastmodels API.
        """
        self.client = Client(api_key=api_key, project_id=project_id)
        self.thread_id = None

    def create_and_run_agent(self, agent_id_initial, messages):
        """
        Create and run agent, send messages and receive a reply.

        :param agent_id_initial: Agent ID.
        :param messages: List of user messages.
        :return: Response content from the agent.
        """
        if self.thread_id is None:
            response = self.client.agent.threads.create_and_run(
                agent_id=agent_id_initial,
                messages=messages
            )
            self.thread_id = response.thread_id
        else:
            response = self.client.agent.threads.create_and_run(
                agent_id=agent_id_initial,
                thread_id=self.thread_id,
                messages=messages
            )

        return response.content[0].text.value

# 使用示例
if __name__ == "__main__":
    config = dict()
    config['API_KEYS']['meeting_minutes_api_key'] = "MGOhqOj79GXdXhShINo3d_g7n36IMoxgWUn2mJsNbBqjWs4odyl2KDUXgxDgxFWNEqnDSvHSEwiTtTVxZuu5yA"
    config['API_KEYS']['meeting_minutes_project_id'] = "6PWZcI42BV4skuwfVlmf5n"
    client = AIClient(config)
    client.create_and_run_agent("1LlvZTbeCoMV4KLQH7X9l5", )

