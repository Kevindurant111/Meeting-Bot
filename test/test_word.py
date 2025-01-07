import os
from docx import Document

def insert_content_after_keyword(file_path, keyword, insert_content):
    """
    在表格中找到包含关键字的行，在该行的下一列插入指定内容。

    :param file_path: Word 文件的路径。
    :param keyword: 要查找的关键字。
    :param insert_content: 要插入的内容。
    :return: 修改后的文档。
    """
    # 加载 Word 文档
    document = Document(file_path)

    # 遍历文档中的所有表格
    for table in document.tables:
        for row in table.rows:
            # 如果该行中有包含关键字的单元格
            if any(keyword in cell.text for cell in row.cells):
                # 在找到的行的第一列后插入内容
                for i, cell in enumerate(row.cells):
                    if keyword in cell.text:
                        # 如果这是目标单元格（会议主题），插入新内容
                        if i + 1 < len(row.cells):
                            row.cells[i + 1].text = insert_content
                        else:
                            # 如果该行没有下一列，则扩展表格并插入新列
                            row.cells[-1].add_paragraph(insert_content)

    # 确保保存目录存在
    output_path = "modified_" + os.path.basename(file_path)
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 保存修改后的文件
    document.save(output_path)
    print(f"修改后的文档已保存为：{output_path}")

# 示例
file_path = "../media/会议纪要-v1.0.1.docx"
insert_content_after_keyword(file_path, "会议主题", "TCB会议")
